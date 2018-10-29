#!/usr/bin/env python3
# -*- coding=utf-8 -*-

# @Time    : 2018-10-18
# @Author  : J.sky
# @Mail    : bosichong@qq.com
# @Site    : www.17python.com
# @Title   : Python实现小学生加减乘除速算考试题卷。
# @Url     : http://www.17python.com/blog/29
# @Details : Python实现小学生加减乘除速算考试题卷。
# @Other   : OS X 10.11.6
#            Python 3.6.1
#            PyCharm
###################################

###################################

import os

from docx import Document  # 引入docx类生成docx文档
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PrintPreview:
    '''本类负责生成完整的口算题文档使之适合打印机打印。可以生成多套题，生成数可以控。

    - @p_list   list
    需要打印口算题库，至少包含一套口算题

    - @p_title   list
    页面标题，这个标题的生成依据程序题型的选择和数字的范围选择而生成，例如：选择了0-20，加减法，进退位
    则自动生成标题为：0到20加减法进退位混合口算题，list中包含了多套题的页面标题名称

    - @p_column  int
    打印页排版口算题的列数

    '''

    p_list = None
    p_title = None
    p_column = None
    p_title_size = None

    def __init__(self, l, tit, col=3, size=26):
        '''
        :param l: list 需要打印的口算题列表
        :param tit: list 口算页标题列表
        :param col: int 列数
        :param size: int 标题字号
        '''
        self.p_list = l
        self.p_title = tit
        self.p_column = col
        self.p_title_size = size

    def create_psmdocx(self, l,title, docxname):
        '''
        :param l list 一组题库
        :param title str 页面标题
        :param docxname  str 题库保存文件名
        :return: none
        '''
        if(title == ''):
            page_title = '小学生口算题'
        else:
            page_title = title
        p_docx = Document()  # 创建一个docx文档
        p_docx.styles['Normal'].font.name = u'Times'  # 可换成word里面任意字体
        p = p_docx.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        run = p.add_run(page_title)
        run.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(self.p_title_size)  # 字体大小设置，和word里面的字号相对应

        # 判断需要用到的行数
        if (len(l) % self.p_column):
            rs = len(l) // self.p_column + 1
        else:
            rs = len(l) // self.p_column

        # print(rows)

        # 将口算题添加到docx表格中
        k = 0  # 计数器
        table = p_docx.add_table(rows=rs, cols=self.p_column)

        for i in range(rs):
            row_cells = table.rows[i].cells
            for j in range(self.p_column):
                if (k > len(l) - 1):
                    break
                else:
                    row_cells[j].text = l[k]
                    k = k + 1
        # table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        table.style.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        table.style.font.size = Pt(16)  # 字体大小设置，和word里面的字号相对应
        p_docx.save('{}.docx'.format(docxname))  # 输出docx

    def produce(self):
        k = 1
        for l,t in zip(self.p_list,self.p_title):
            self.create_psmdocx(l,t,t+str(k))
            k = k+1





if __name__ == '__main__':
    l = [['1-17=', '3-4=', '13-6=', '15-5=', '2-4=', '15-9=', '12-13=', '15-12=', '14-16=', '4-11=', '18-16=', '12-14=',
         '14-7=', '7-17=', '16-19=', '9-4=', '13-3=', '7-4=', '11-1=', '4-5=', '16-18=', '11-6=', '5-11=', '12-6=',
         '4-6=', '5-2=', '8-12=', '14-15=', '15-19=', '1-12=', '7-13=', '6-4=', '6-12=', '18-12=', '11-15=', '18-17=',
         '1-16=', '3-12=', '14-2=', '16-4=', '7-16=', '6-19=', '16-9=', '8-2=', '13-9=', '17-6=', '6-7=', '17-5=',
         '9-14=', '4-2=', '3-5=', '12-4=', '12-15=', '8-5=', '19-12=', '1-5=', '14-19=', '15-4=', '8-18=', '18-13=',
         '8-11=', '1-7=', '12-8=', '2-3=', '18-15=', '7-5=', '5-16=', '6-3=', '7-12=', '3-16=', '12-18=', '16-11=',
         '6-18=', '12-2=', '1-4=', '11-13=', '4-19=', '16-15=', '17-2=', '7-14='],['1-17=', '3-4=', '13-6=', '15-5=', '2-4=', '15-9=', '12-13=', '15-12=', '14-16=', '4-11=', '18-16=', '12-14=',
         '14-7=', '7-17=', '16-19=', '9-4=', '13-3=', '7-4=', '11-1=', '4-5=', '16-18=', '11-6=', '5-11=', '12-6=',
         '4-6=', '5-2=', '8-12=', '14-15=', '15-19=', '1-12=', '7-13=', '6-4=', '6-12=', '18-12=', '11-15=', '18-17=',
         '1-16=', '3-12=', '14-2=', '16-4=', '7-16=', '6-19=', '16-9=', '8-2=', '13-9=', '17-6=', '6-7=', '17-5=',
         '9-14=', '4-2=', '3-5=', '12-4=', '12-15=', '8-5=', '19-12=', '1-5=', '14-19=', '15-4=', '8-18=', '18-13=',
         '8-11=', '1-7=', '12-8=', '2-3=', '18-15=', '7-5=', '5-16=', '6-3=', '7-12=', '3-16=', '12-18=', '16-11=',
         '6-18=', '12-2=', '1-4=', '11-13=', '4-19=', '16-15=', '17-2=', '7-14=']]
    t = ['小学生口算题','小学生口算题']
    pp = PrintPreview(l, t, 4)
    pp.produce()